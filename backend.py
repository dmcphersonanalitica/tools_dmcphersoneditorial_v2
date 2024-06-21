import os
import torch
import tempfile
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


quality_model = {
        1: 'whisper-tiny',
        2: 'whisper-base',
        3: 'whisper-small',
        4: 'whisper-medium',
        5: 'whisper-large-v2',
        6: 'whisper-large-v3'
    }


def audio_transcription(quality, audio_file, format_result):
    device = "cuda:0" if torch.cuda.is_available() else "cpu"
    torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

    model_id = "openai/" + quality_model[quality]

    model = AutoModelForSpeechSeq2Seq.from_pretrained(
        model_id, torch_dtype=torch_dtype, low_cpu_mem_usage=True, use_safetensors=True
    )
    model.to(device)

    processor = AutoProcessor.from_pretrained(model_id)

    pipe = pipeline(
        "automatic-speech-recognition",
        model=model,
        tokenizer=processor.tokenizer,
        feature_extractor=processor.feature_extractor,
        max_new_tokens=128,
        chunk_length_s=25,
        batch_size=16,
        torch_dtype=torch_dtype,
        device=device,
    )

    # dataset = load_dataset("distil-whisper/librispeech_long", "clean", split="validation")
    # sample = dataset[0]["audio"]

    if format_result == 'Timeline text':
        result = pipe(audio_file, return_timestamps=True)
        # full_result = ''
        # segments = result["segments"]
        # for seg in segments:
        #    full_result += str(seg['start']) + '   --->   ' + str(seg['end']) + '\n\n' + seg['text'] + '\n\n'

        # return full_result
        return result["chunks"]

    result = pipe(audio_file)

    return result["text"]


def path_file(file):
    temp_dir = tempfile.mkdtemp()
    path = os.path.join(temp_dir, file.name)
    with open(path, "wb") as f:
        f.write(file.getvalue())
    return path.replace('\\', '\\\\')


def save_to_word(filename, text, format):
    document = Document()

    header = document.add_heading('DMcPherson Editorial Tools\n')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(text)

    if format == 'Timeline text':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(filename.name + '.docx')
    return document
